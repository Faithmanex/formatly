"""
Spacing removal utility for cleaning document paragraph spacing.
"""
from docx.shared import Pt
from docx.oxml.ns import qn


def _remove_para_spacing(paragraph):
    """Internal helper to reset space_before and space_after for a paragraph."""
    # 1. Remove the old spacing tag first (clears all: before, after, line spacing)
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        pPr.remove(spacing)

    # 2. Apply fresh zero spacing
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def remove_all_spacing(doc):
    """
    Remove all space_before and space_after from every paragraph in the document,
    including those in tables, headers, and footers.
    """
    # 1. Main body paragraphs
    for paragraph in doc.paragraphs:
        _remove_para_spacing(paragraph)

    # 2. Table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _remove_para_spacing(paragraph)

    # 3. Headers and footers in all sections
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    _remove_para_spacing(paragraph)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    _remove_para_spacing(paragraph)

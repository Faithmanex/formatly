"""
Formatly Desktop Core Formatter
-------------------------------
This module contains the core formatting logic specifically for the Desktop application.

MAINTENANCE WARNING:
    This file is currently a duplicate of the root `core/formatter.py`.
    It exists here to keep the desktop application self-contained.
    Any changes made to the logic here should likely be mirrored in the root `core/`
    and vice-versa, until a shared library structure is fully implemented.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from .style_guides import STYLE_GUIDES
from core.api_clients import AIClient
from core.constants import STYLE_MAPPINGS
from utils.spacing import remove_all_spacing
import json
import os
import re
import traceback
import json5
import time
import csv
from datetime import datetime
from pathlib import Path

class DocumentStructureManager:
    """Manages document structure initialization, validation, and prompt creation."""
    
    @staticmethod
    def get_default_structure() -> dict:
        """Returns the default document structure template."""
        return {
            "blocks": [
                {
                    "type": "", 
                    "content": "",
                    "attributes": {}
                }
            ]
        }
    
    @staticmethod
    def create_detection_prompt(paragraphs: list, english_variant: str = "us") -> tuple[str, str]:
        """Creates system and user prompts for document structure detection."""

        system_prompt = """\
You are an expert academic document parser. Your task is to read a full academic document \
and classify every paragraph into a structured JSON block list.

═══════════════════════════════════════════════════════
STEP 1 — UNDERSTAND THE DOCUMENT BEFORE CLASSIFYING
═══════════════════════════════════════════════════════
Before you assign a single block type, read the entire input and build an internal model of:
  • What kind of document this is (thesis, dissertation, journal article, research paper, report, etc.)
  • The academic field and discipline (affects heading conventions, reference style, table/figure norms)
  • The overall structure: how many chapters, whether there is a title page, abstract, front matter, appendices
  • The heading hierarchy actually used by this author (some use "CHAPTER X + Title", some use "1. Title", some use bare titles)
  • Recurring patterns: how the author formats lists, captions, block quotes, footnotes
  • Any unusual conventions specific to this document

Use this understanding to inform every classification decision. A line that looks ambiguous in isolation will often be unambiguous when you know the document's structure and conventions.

═══════════════════════════════════════════════════════
STEP 2 — EXTRACTION RULES (non-negotiable)
═══════════════════════════════════════════════════════
1. VERBATIM TEXT — Copy content exactly as it appears. Do NOT fix typos, reword, reorder, remove words, or alter spacing. Extract every character as-is.

2. COMPLETENESS — Every non-empty line of input must become a block. Nothing may be omitted.

3. NO EMPTY BLOCKS — Never output a block with content: "". Skip blank/empty lines entirely.

4. ONE BLOCK = ONE SEMANTIC ELEMENT — This is the most critical structural rule:
   Each block must represent exactly one distinct semantic unit.
   - If a single line contains two different types of information, split it into two blocks.
   - Never use "\\n" inside content to merge two semantically different elements.
   - Exception: a naturally multi-line element of the same type (e.g. a multi-sentence abstract paragraph, or a multi-line postal address for a single institution) may use "\\n".

   CORRECT:
     {{"type": "title_department", "content": "Department of Linguistics", "attributes": {{}}}}
     {{"type": "institution",      "content": "University of Lagos",        "attributes": {{}}}}

   WRONG:
     {{"type": "institution", "content": "Department of Linguistics\\nUniversity of Lagos", "attributes": {{}}}}

5. REQUIRED FIELDS — Every block must have exactly these three fields, no more, no less:
   {{"type": "...", "content": "...", "attributes": {{}}}}

6. VALID JSON — Output must be strictly valid, parseable JSON.

7. BLOCK BOUNDARIES — Each block is a complete JSON object: opens with {{ closes with }}, separated by commas. A single object must never contain two "type" keys.

═══════════════════════════════════════════════════════
STEP 3 — HEADING CLASSIFICATION
═══════════════════════════════════════════════════════
Use the document's own heading conventions (identified in Step 1) to guide level assignment.

CHAPTER LABELS
  "CHAPTER ONE", "CHAPTER 1", "CHAPTER I", "CHAPTER IV" → always heading_1.
  These are standalone label lines. The title that follows on the next line is also heading_1.
  Do NOT merge them — output each as a separate heading_1 block.

NUMBERED SECTIONS
  Top-level numbers (1., 2., 1.0, I., II.) → heading_1
  One level of subdivision (1.1, 2.3, I.A) → heading_2
  Two levels of subdivision (1.1.1, 2.3.4) → heading_3
  And so on for deeper levels.

UNNUMBERED HEADINGS
  Use the document's positional and contextual evidence:
  - Where does it appear relative to surrounding content?
  - Does it introduce a major section or a sub-section?
  - Is it consistent with other headings at the same visual level in this document?
  Classify based on your understanding of the document's structure — not on surface features alone.

PRELIMINARY PAGE HEADINGS
  Use the specific type if available (abstract_heading, acknowledgements_heading, etc.).
  Fall back to heading_1 only if no specific type applies.

═══════════════════════════════════════════════════════
STEP 4 — LIST CLASSIFICATION
═══════════════════════════════════════════════════════
Every list item block must include a "list_id" integer inside "attributes".

  list_item_bullet   → unordered list (bullet points, dashes, dots)
  list_item_number   → ordered list with numerals (1. 2. 3.)
  list_item_alphabet → ordered list with letters (a. b. c. or A. B. C.)

GROUPING: Items belonging to the same list share the same list_id.
RESTART: When a new list begins after intervening non-list content, assign a new list_id.

Examples:
  {{"type": "list_item_bullet",   "content": "Qualitative analysis",  "attributes": {{"list_id": 1}}}}
  {{"type": "list_item_number",   "content": "Collect data",          "attributes": {{"list_id": 2}}}}
  {{"type": "list_item_alphabet", "content": "Review literature",     "attributes": {{"list_id": 3}}}}

═══════════════════════════════════════════════════════
STEP 5 — APPENDIX CLASSIFICATION
═══════════════════════════════════════════════════════
  appendices_heading → the master title standing alone: "APPENDICES", "LIST OF APPENDICES"
  appendix_heading   → individual appendix titles: "Appendix A", "Appendix 1: Survey Instrument"

═══════════════════════════════════════════════════════
STEP 6 — FULL BLOCK TYPE REFERENCE
═══════════════════════════════════════════════════════
TITLE PAGE ELEMENTS
  title               — Main title of the work
  title_byline        — Subtitle or byline beneath the title
  author              — Author name(s)
  institution         — University or organisation name
  title_department    — Department or faculty name
  registration_number — Student/registration ID
  degree              — Degree being submitted for
  course              — Course name or code
  instructor          — Supervisor, lecturer, or instructor name
  due_date            — Submission or due date

FRONT MATTER
  abstract_heading         — The word "Abstract" as a section heading
  abstract_text            — Abstract body paragraph(s)
  keywords                 — Keywords line (e.g. "Keywords: X, Y, Z")
  epigraph                 — Epigraph or opening quote
  dedication_heading       — The word "Dedication"
  dedication_body          — Dedication text
  acknowledgements_heading — The word "Acknowledgements" / "Acknowledgments"
  acknowledgement_body     — Acknowledgements body paragraph(s)
  preface_heading          — The word "Preface"
  preface_body             — Preface body paragraph(s)
  table_of_contents_heading — Title like "Table of Contents" or "Contents"
  table_of_contents_item    — An entry in the table of contents

HEADINGS
  heading_1 — Top-level chapter or section heading
  heading_2 — Sub-section heading
  heading_3 — Sub-sub-section heading
  heading_4 — Fourth-level heading
  heading_5 — Fifth-level heading

BODY CONTENT
  body        — Standard prose paragraph
  block_quote — Indented or set-off quotation

LISTS (always include list_id in attributes)
  list_item_number   — Numbered list item
  list_item_bullet   — Bullet list item
  list_item_alphabet — Alphabetical list item

TABLES & FIGURES
  table_caption — Caption above or below a table
  table_note    — Note beneath a table (e.g. "Note. ...")
  table_data    — A row or cell of tabular data presented as text
  figure_caption — Caption for a figure, chart, or image

REFERENCES
  references_heading  — The heading "References", "Bibliography", "Works Cited", etc.
  reference_list_item — A single formatted reference entry

APPENDICES
  appendices_heading — Master appendix section title
  appendix_heading   — Individual appendix title

OTHER
  code_block — Code listing or verbatim technical content
  footnote   — Footnote text

═══════════════════════════════════════════════════════
OUTPUT FORMAT
═══════════════════════════════════════════════════════
Return a single JSON object:
{{
    "blocks": [
        {{"type": "...", "content": "...", "attributes": {{}}}},
        ...
    ]
}}

No markdown fences. No preamble. No explanation. JSON only.\
"""

        user_prompt = (
            f"Document content:\n\n{chr(10).join(paragraphs)}\n\n"
            "Classify every paragraph and return JSON only."
        )

        return system_prompt, user_prompt
    
    @staticmethod
    def parse_and_validate_response(response_text: str) -> dict:
        """Parses and validates AI response into document structure."""
        text = response_text.strip()
        
        for marker in ["```json", "```"]:
            if text.startswith(marker):
                text = text[len(marker):]
            if text.endswith("```"):
                text = text[:-3]
        
        text = text.strip()
        
        if not text or text[0] not in ['{', '[']:
            raise json.JSONDecodeError(f"Invalid JSON: {text[:100]}...", text, 0)
        
        try:
            data = json5.loads(text)
            return DocumentStructureManager.validate_structure(data)
        except (json.JSONDecodeError, ValueError) as e:
            print(f"JSON parse error: {e}")
            print(f"Raw response snippet: {response_text[:500]}...")
            raise
    
    @staticmethod
    def validate_structure(data: dict) -> dict:
        """Validates and sanitizes document structure."""
        if not isinstance(data, dict):
            raise ValueError("Document structure must be a dictionary")
        
        template = DocumentStructureManager.get_default_structure()
        
        def deep_merge(template_dict, data_dict):
            result = template_dict.copy()
            for key, value in data_dict.items():
                if key in result and isinstance(result[key], dict) and isinstance(value, dict):
                    result[key] = deep_merge(result[key], value)
                else:
                    result[key] = value
            return result
        
        return deep_merge(template, data)

class AdvancedFormatter:
    def __init__(self, style_name: str, ai_client: AIClient, english_variant: str = "us"):
        self.selected_style_name = style_name.lower()
        self.selected_style_guide = STYLE_GUIDES.get(self.selected_style_name, STYLE_GUIDES["apa"])
        self.ai_client = ai_client
        self.english_variant = english_variant
        
        print(f"Loaded Style: {self.selected_style_name} | Language: {self.english_variant}")
        self._compile_reference_patterns()
        self._compile_inline_patterns()
        self._doc_styles = None

    def _compile_reference_patterns(self):
        """Pre-compile reference formatting patterns for efficiency."""
        ref_config = self.selected_style_guide.get('reference_formatting', {})
        self.compiled_ref_patterns = []
        for pattern in ref_config.get('patterns', []):
            if 'regex' in pattern and 'formatting' in pattern:
                self.compiled_ref_patterns.append({
                    'pattern': re.compile(pattern['regex'], re.DOTALL),
                    'formatting': pattern['formatting']
                })

    def _compile_inline_patterns(self):
        """Pre-compile inline formatting patterns."""
        inline_config = self.selected_style_guide.get('inline_formatting', [])
        self.compiled_inline_patterns = []
        for rule in inline_config:
            if 'regex' in rule and 'formatting' in rule:
                self.compiled_inline_patterns.append({
                    'pattern': re.compile(rule['regex']),
                    'formatting': rule['formatting'],
                    'target_styles': rule.get('target_styles', [])
                })

    def _apply_inline_formatting(self, doc):
        """
        Apply inline formatting based on regex patterns defined in style guide.
        """
        if not hasattr(self, 'compiled_inline_patterns') or not self.compiled_inline_patterns:
            return
        
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            style_name = p.style.name if hasattr(p.style, 'name') else str(p.style)
            
            for rule in self.compiled_inline_patterns:
                if rule['target_styles'] and style_name not in rule['target_styles']:
                    continue
                for match in rule['pattern'].finditer(p.text):
                    start_idx, end_idx = match.span()
                    self._apply_formatting_to_range(p, start_idx, end_idx, rule['formatting'])

    def _apply_formatting_to_range(self, paragraph, start_char, end_char, formatting):
        self._ensure_run_boundary(paragraph, start_char)
        self._ensure_run_boundary(paragraph, end_char)
        current_pos = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            run_end = current_pos + run_len
            if current_pos >= start_char and run_end <= end_char:
                if not run.font: run.font = paragraph.style.font 
                if 'italic' in formatting: run.font.italic = formatting['italic']
                if 'bold' in formatting: run.font.bold = formatting['bold']
            current_pos = run_end

    def _ensure_run_boundary(self, paragraph, char_index):
        current_pos = 0
        for i, run in enumerate(paragraph.runs):
            run_text = run.text
            run_len = len(run_text)
            if current_pos < char_index < current_pos + run_len:
                split_point = char_index - current_pos
                part1 = run_text[:split_point]
                part2 = run_text[split_point:]
                run.text = part1
                new_run = paragraph.add_run(part2)
                self._copy_run_formatting(run, new_run)
                p_element = paragraph._p
                run_element = run._element
                new_run_element = new_run._element
                p_element.remove(new_run_element)
                run_element.addnext(new_run_element)
                return
            current_pos += run_len

    def _get_style(self, doc, style_name):
        """Get an existing style from the document."""
        return doc.styles[style_name]

    def _log_stats(self, input_path, duration, usage_stats):
        """Log formatting stats to console."""
        prompt_tokens = 0
        completion_tokens = 0
        total_tokens = 0
        
        if usage_stats:
            if hasattr(usage_stats, 'prompt_tokens'):
                prompt_tokens = usage_stats.prompt_tokens
                completion_tokens = usage_stats.completion_tokens
                total_tokens = usage_stats.total_tokens
            else:
                print("Warning: usage_stats object has unexpected shape; token counts unavailable.")
        
        print("-" * 30)
        print(f"⏱️  Duration:      {duration:.2f}s")
        print(f"🔢 Total Tokens:  {total_tokens}")
        print(f"   • Prompt:      {prompt_tokens}")
        print(f"   • Completion:  {completion_tokens}")
        print("-" * 30)

    def format_document(self, input_path, output_path, doc=None):
        if doc is None:
            doc = Document(input_path)
        
        self._doc_styles = doc.styles

        self._remove_leading_whitespace(doc)
        self._clear_run_formatting(doc)
        remove_all_spacing(doc)

        self._customize_builtin_styles(doc)
        self._apply_margins(doc)

        start_time = time.time()

        paragraphs_text = [p.text for p in doc.paragraphs]
        doc_structure, usage_stats = self._detect_paragraph_types(paragraphs_text)

        has_title_page, title_end_idx, body_start_idx = self._detect_and_manage_title_page(doc, doc_structure)
        self._add_page_numbers(doc, has_title_page, title_end_idx, body_start_idx)

        self._deduplicate_consecutive_headings(doc, doc_structure)
        self._join_headings(doc, doc_structure)

        self._format_content_in_place(doc, doc_structure, has_title_page, title_page_boundary)

        self._format_references(doc, doc_structure)

        self._apply_font_properties(doc)
        self._format_tables(doc)

        self._apply_explicit_paragraph_properties(doc)
        self._apply_list_properties(doc)

        self._apply_heading_styles(doc)
        
        self._apply_inline_formatting(doc)

        self._remove_blank_lines(doc, title_page_boundary)

        output_path_obj = Path(output_path)
        final_output_path = str(output_path_obj)
        
        for i in range(10):
            try:
                doc.save(final_output_path)
                break
            except PermissionError:
                new_stem = f"{output_path_obj.stem} ({i + 1})"
                final_output_path = str(output_path_obj.with_name(f"{new_stem}{output_path_obj.suffix}"))
                print(f"Permission denied. Retrying with: {final_output_path}")
            except Exception as e:
                raise e
        else:
            raise PermissionError(f"Could not save document after 10 attempts. Last path tried: {final_output_path}")
        
        end_time = time.time()
        duration = end_time - start_time
        self._log_stats(input_path, duration, usage_stats)
        
        return final_output_path

    def _customize_builtin_styles(self, doc):
        """
        Customize built-in styles and add custom styles based on the style guide.
        """
        styles_config = self.selected_style_guide["styles"]
        doc_styles = doc.styles
        
        for style_name, style_config in styles_config.items():
            try:
                if style_name not in doc_styles:
                    style_type = style_config.get("type", WD_STYLE_TYPE.PARAGRAPH)
                    doc.styles.add_style(style_name, style_type)

                style = self._get_style(doc, style_name)
                
                if "based_on" in style_config and style_config["based_on"] in doc_styles:
                    style.base_style = doc_styles[style_config["based_on"]]
                
                if "font" in style_config:
                    font_config = style_config["font"]
                    font = style.font

                    if "name" in font_config:
                        font.name = font_config["name"]
                    if "size" in font_config:
                        font.size = font_config["size"]
                    if "bold" in font_config:
                        font.bold = font_config["bold"]
                    if "italic" in font_config:
                        font.italic = font_config["italic"]
                    if "underline" in font_config:
                        font.underline = font_config["underline"]
                    if "color" in font_config and font_config["color"]:
                        font.color.rgb = font_config["color"]
                    if "all_caps" in font_config:
                        font.all_caps = font_config["all_caps"]

                if "paragraph" in style_config:
                    para_format = style.paragraph_format
                    para_config = style_config["paragraph"]
                    
                    if "name" in para_config:
                        para_format.name = para_config["name"]
                    if "alignment" in para_config:
                        para_format.alignment = para_config["alignment"]
                    if "left_indent" in para_config:
                        para_format.left_indent = para_config["left_indent"]
                    if "right_indent" in para_config:
                        para_format.right_indent = para_config["right_indent"]
                    if "first_line_indent" in para_config:
                        para_format.first_line_indent = para_config["first_line_indent"]
                    if "space_before" in para_config:
                        para_format.space_before = para_config["space_before"]
                    if "space_after" in para_config:
                        para_format.space_after = para_config["space_after"]
                    if "line_spacing" in para_config:
                        para_format.line_spacing_rule = para_config["line_spacing"]
                    if "keep_together" in para_config:
                        para_format.keep_together = para_config["keep_together"]
                    if "keep_with_next" in para_config:
                        para_format.keep_with_next = para_config["keep_with_next"]
                    if "page_break_before" in para_config:
                        para_format.page_break_before = para_config["page_break_before"]
                    if "widow_control" in para_config:
                        para_format.widow_control = para_config["widow_control"]
                    if "orphan_control" in para_config:
                        try:
                            para_format.orphan_control = para_config["orphan_control"]
                        except AttributeError:
                            pass
                    if "outline_level" in para_config:
                        try:
                            para_format.outline_level = para_config["outline_level"]
                        except AttributeError:
                            pass
                
                if hasattr(style, 'hidden'):
                    style.hidden = style_config.get("hidden", False)
                if hasattr(style, 'unhide_when_used'):
                    style.unhide_when_used = style_config.get("unhide_when_used", True)
                if hasattr(style, 'quick_style'):
                    style.quick_style = style_config.get("quick_style", True)
                
                if (style_name.startswith("Heading") and " " in style_name) or style_name == "Appendices Title":
                    try:
                        level = int(style_name.split(' ')[1]) if "Heading" in style_name else 1
                        if hasattr(style.paragraph_format, 'outline_level'):
                            style.paragraph_format.outline_level = level - 1
                        
                        pPr = style._element.get_or_add_pPr()
                        outlineLvl = pPr.find(qn('w:outlineLvl'))
                        if outlineLvl is None:
                            outlineLvl = OxmlElement('w:outlineLvl')
                            pPr.append(outlineLvl)
                        outlineLvl.set(qn('w:val'), str(level - 1))
                        
                    except (ValueError, IndexError, AttributeError) as e:
                        print(f"Warning: Failed to set outline level for {style_name}: {e}")
                        pass
                
                if "next_style" in style_config and style_config["next_style"] in doc_styles:
                    style.next_paragraph_style = doc_styles[style_config["next_style"]]
                
                self._remove_borders(style)

            except Exception as e:
                print(f"Warning: Could not apply style '{style_name}': {str(e)}")
                continue

    def _remove_borders(self, style):
        """Removes all borders (bottom, top, left, right) from a paragraph style."""
        pPr = style._element.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)

    def _apply_margins(self, doc):
        section = doc.sections[0]
        margin_map = {
            "left": "left_margin", "right": "right_margin",
            "top": "top_margin", "bottom": "bottom_margin",
            "header": "header_distance", "footer": "footer_distance",
            "gutter": "gutter"
        }
        for key, value in self.selected_style_guide["margins"].items():
            if key in margin_map:
                setattr(section, margin_map[key], value)

    def _format_tables(self, doc):
        """
        Format all tables in the document.
        Ensures cells use the 'Table Content' style to prevent text cutoff
        caused by the 'Normal' style's indent.
        """
        if "Table Content" not in doc.styles:
            return

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.style.name.startswith("Heading"):
                            continue
                        paragraph.style = doc.styles["Table Content"]

    def _remove_leading_whitespace(self, doc):
        """
        Strip leading and trailing spaces/tabs from every paragraph.

        Operates at the run level so that inline formatting (bold, italic, etc.)
        on the remaining characters is fully preserved.

        Leading pass  — walks runs left-to-right; clears any all-whitespace run
                        before the first content run, then lstrips the first
                        content run.
        Trailing pass — walks runs right-to-left; clears any all-whitespace run
                        after the last content run, then rstrips the last
                        content run.
        """
        for paragraph in doc.paragraphs:
            runs = paragraph.runs
            if not runs:
                continue

            # --- Leading strip ---
            for run in runs:
                text = run.text
                if not text:
                    continue
                if text.strip() == "":
                    run.text = ""
                else:
                    run.text = text.lstrip()
                    break

            # --- Trailing strip ---
            for run in reversed(runs):
                text = run.text
                if not text:
                    continue
                if text.strip() == "":
                    run.text = ""
                else:
                    run.text = text.rstrip()
                    break

    def _clear_run_formatting(self, doc):
        """Clear old run formatting (character styles, fonts, colors, effects) from all paragraphs (body and tables)."""
        
        def clear_para_runs(paragraph):
            for run in paragraph.runs:
                run.bold = None
                run.italic = None
                run.underline = None
                
                if hasattr(run, 'font'):
                    run.font.name = None
                    run.font.size = None
                    try:
                        if run.font.color:
                            run.font.color.rgb = None
                    except (ValueError, AttributeError):
                        pass

        for paragraph in doc.paragraphs:
            clear_para_runs(paragraph)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        clear_para_runs(paragraph)

    def _detect_paragraph_types(self, paragraphs):
        system_prompt, user_prompt = DocumentStructureManager.create_detection_prompt(paragraphs, self.english_variant)
        
        max_retries = 3
        retry_count = 0
        
        while retry_count < max_retries:
            try:
                full_text, usage_stats = self.ai_client.detect_structure(system_prompt, user_prompt)
                
                if not full_text:
                    raise ValueError("API returned an empty response. Please try again.")
                
                return DocumentStructureManager.parse_and_validate_response(full_text), usage_stats

            except (json.JSONDecodeError, ValueError) as e:
                retry_count += 1
                error_msg = f"JSON Parsing Error: {str(e)}"
                print(f"ERROR: {error_msg}")
                
                if retry_count < max_retries:
                    wait_time = 2 * retry_count
                    print(f"⚠️ Parsing failed. Retrying in {wait_time}s (Attempt {retry_count}/{max_retries})...")
                    time.sleep(wait_time)
                else:
                    print(f"Response parsing failed. The API response may be incomplete or malformed.")
                    raise ValueError(error_msg) from e
            except Exception as e:
                error_msg = f"API Request Error: {str(e)}"
                print(f"ERROR: {error_msg}")
                traceback.print_exc()
                raise ValueError(error_msg) from e

    def _add_page_numbers(self, doc, has_title_page, title_end_idx, body_start_idx):
        from copy import deepcopy
        
        meta = self.selected_style_guide["meta"]
        page_numbering = meta.get("page_numbering")
        position = meta.get("page_numbers") if not page_numbering else page_numbering.get("body_position")
        
        if not position and not page_numbering:
            return

        if has_title_page:
            # Step A: Split Section 1 (Title) and Section 2 (Front Matter)
            if title_end_idx > 0 and title_end_idx < len(doc.paragraphs):
                p_fm_start = doc.paragraphs[title_end_idx - 1]
                pPr = p_fm_start._p.get_or_add_pPr()
                if not pPr.xpath('w:sectPr'):
                    pPr.append(deepcopy(doc.sections[0]._sectPr))

            # Step B: Split Section 2 (Front Matter) and Section 3 (Body)
            if body_start_idx > title_end_idx and body_start_idx < len(doc.paragraphs):
                p_body_start = doc.paragraphs[body_start_idx - 1]
                pPr = p_body_start._p.get_or_add_pPr()
                if not pPr.xpath('w:sectPr'):
                    pPr.append(deepcopy(doc.sections[-1]._sectPr))

            title_section = doc.sections[0]
            title_section.different_first_page_header_footer = True
            for container in [title_section.first_page_header, title_section.first_page_footer]:
                for p in list(container.paragraphs):
                    p._element.getparent().remove(p._element)

            if len(doc.sections) >= 3:
                fm_section = doc.sections[1]
                body_section = doc.sections[2]
                
                if page_numbering:
                    fm_format = page_numbering.get("front_matter_format", "decimal")
                    pgNumType_fm = fm_section._sectPr.find(qn('w:pgNumType'))
                    if pgNumType_fm is None:
                        pgNumType_fm = OxmlElement('w:pgNumType')
                        fm_section._sectPr.append(pgNumType_fm)
                    pgNumType_fm.set(qn('w:fmt'), 'lowerRoman' if fm_format == 'roman' else ('decimal' if fm_format == 'arabic' else fm_format))

                pgNumType_body = body_section._sectPr.find(qn('w:pgNumType'))
                if pgNumType_body is None:
                    pgNumType_body = OxmlElement('w:pgNumType')
                    body_section._sectPr.append(pgNumType_body)
                pgNumType_body.set(qn('w:start'), '1')
                body_fmt = page_numbering.get("body_format", "decimal") if page_numbering else "decimal"
                pgNumType_body.set(qn('w:fmt'), 'decimal' if body_fmt == 'arabic' else body_fmt)
                
                section_for_numbering = body_section
                section_for_fm = fm_section
                
            elif len(doc.sections) == 2: 
                body_section = doc.sections[1]
                pgNumType = body_section._sectPr.find(qn('w:pgNumType'))
                if pgNumType is None:
                    pgNumType = OxmlElement('w:pgNumType')
                    body_section._sectPr.append(pgNumType)
                pgNumType.set(qn('w:start'), '1')
                pgNumType.set(qn('w:fmt'), 'decimal')
                section_for_numbering = body_section
                section_for_fm = None
            else:
                section_for_numbering = doc.sections[0]
                section_for_fm = None
        else:
            section_for_numbering = doc.sections[0]
            section_for_fm = None

        if has_title_page and page_numbering and section_for_fm:
             fm_pos = page_numbering.get("front_matter_position", "center")
             container1 = section_for_fm.header if "header" in fm_pos else section_for_fm.footer
             p1 = container1.paragraphs[0] if container1.paragraphs else container1.add_paragraph()
             p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT if "right" in fm_pos else (WD_ALIGN_PARAGRAPH.CENTER if "center" in fm_pos else WD_ALIGN_PARAGRAPH.LEFT)
             self._insert_page_field_to_paragraph(p1)

        container = section_for_numbering.header if "header" in position else section_for_numbering.footer
        p_body = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        p_body.alignment = WD_ALIGN_PARAGRAPH.RIGHT if "right" in position else (WD_ALIGN_PARAGRAPH.CENTER if "center" in position else WD_ALIGN_PARAGRAPH.LEFT)
        self._insert_page_field_to_paragraph(p_body)

    def _insert_page_field_to_paragraph(self, paragraph):
         run = paragraph.add_run()
         fldChar_begin = OxmlElement('w:fldChar')
         fldChar_begin.set(qn('w:fldCharType'), 'begin')
         instrText = OxmlElement('w:instrText')
         instrText.set(qn('xml:space'), 'preserve')
         instrText.text = "PAGE"
         fldChar_end = OxmlElement('w:fldChar')
         fldChar_end.set(qn('w:fldCharType'), 'end')
         run._r.extend([fldChar_begin, instrText, fldChar_end])

    def _detect_and_manage_title_page(self, doc, doc_structure):
        """
        Detects Title Page & Front matter breakout dynamically by scanning document block types.
        Returns has_title_page (bool), title_end_idx (int), body_start_idx (int).
        """
        blocks = doc_structure.get("blocks", [])
        if not blocks:
            return False, 0, 0
            
        title_styles = {
            "title", "author", "institution", "course", "instructor", "due_date", 
            "registration_number", "degree", "title_department", "title_byline"
        }
        
        has_title = any(b.get("type") in title_styles for b in blocks[:10])
        if not has_title:
             return False, 0, 0
             
        # Heuristic 1: Find end of Title Page (Last title element position)
        last_title_idx = -1
        for i, b in enumerate(blocks):
             if b.get("type") in title_styles:
                  last_title_idx = i
        
        title_end_content = blocks[last_title_idx].get("content", "").strip() if last_title_idx != -1 else None
             
        # Heuristic 2: Find start of body: usually first heading_1 that is not front matter
        body_start_content = None
        for b in blocks:
             t = b.get("type")
             content = b.get("content", "").lower()
             if t == "heading_1" and not any(f in content for f in ["dedication", "preface", "contents", "abstract"]):
                  body_start_content = b.get("content", "").strip()
                  break
                  
        title_end_idx = 0
        body_start_idx = len(doc.paragraphs) if body_start_content else 0
        for idx, p in enumerate(doc.paragraphs):
             text = p.text.strip()
             if title_end_content and text == title_end_content:
                  title_end_idx = idx + 1 # Endpoint is paragraph AFTER the item
             if body_start_content and text == body_start_content:
                  body_start_idx = idx
                  break
                  
        return True, title_end_idx, body_start_idx

    # -------------------------------------------------------------------------
    # Heading deduplication
    # -------------------------------------------------------------------------

    def _deduplicate_consecutive_headings(self, doc, doc_structure):
        """
        Remove consecutive heading blocks that share identical content.

        When the AI returns two (or more) adjacent heading_1 blocks with the
        same text it means the writer accidentally duplicated the line.  The
        second block is noise and must be removed before _join_headings runs,
        otherwise the joiner would merge the duplicate into the label instead
        of the real title.

        Both the XML paragraph and the block list entry are removed so that
        downstream steps see a clean, consistent state.
        """
        blocks = doc_structure.get("blocks", [])
        if not blocks:
            return

        # Build lookup: stripped text → list of paragraph objects (order preserved)
        text_to_paras: dict[str, list] = {}
        for p in doc.paragraphs:
            key = p.text.strip()
            if key:
                text_to_paras.setdefault(key, []).append(p)

        i = 0
        while i < len(blocks) - 1:
            current = blocks[i]
            nxt     = blocks[i + 1]

            current_type    = current.get("type", "")
            current_content = current.get("content", "").strip()
            next_type       = nxt.get("type", "")
            next_content    = nxt.get("content", "").strip()

            is_heading = current_type.startswith("heading") or current_type in {
                "appendix_heading", "appendices_heading",
                "abstract_heading", "acknowledgements_heading",
                "dedication_heading", "preface_heading", "references_heading",
                "table_of_contents_heading",
            }

            if (
                is_heading
                and current_type == next_type
                and current_content == next_content
                and current_content  # guard against empty blocks
            ):
                # Remove the duplicate paragraph from the XML tree
                paras = text_to_paras.get(next_content, [])
                if len(paras) >= 2:
                    dup_p = paras[1]
                    parent = dup_p._element.getparent()
                    if parent is not None:
                        parent.remove(dup_p._element)
                    text_to_paras[next_content] = [paras[0]] + paras[2:]
                    print(f"Removed duplicate heading: '{next_content}'")

                # Remove the duplicate block regardless of whether we found a
                # matching paragraph (keeps block list and document in sync)
                blocks.pop(i + 1)
                # Do NOT increment i — re-check this position for triple duplicates
            else:
                i += 1

    # -------------------------------------------------------------------------
    # Heading join
    # -------------------------------------------------------------------------

    def _build_chapter_separator(self) -> str:
        """
        Returns the separator to place between a chapter label and its title
        when the AI has indicated they belong together (two consecutive heading_1
        blocks where the first is the label and the second is the title).
        """
        style_separators = {
            "apa":      ": ",
            "mla":      ": ",
            "chicago":  ": ",
            "ieee":     ": ",
            "harvard":  " \u2013 ",
            "ama":      ": ",
        }
        return style_separators.get(self.selected_style_name, ": ")

    def _join_headings(self, doc, doc_structure):
        """
        Merges pairs of consecutive heading_1 blocks that the AI has placed on
        separate lines but that belong together as a single heading (e.g. a bare
        "CHAPTER 1" label followed immediately by the chapter title).

        The AI is the sole authority on which blocks should be joined — this
        method does not apply any additional heuristics.  The contract is:

            Two consecutive heading_1 blocks (no non-empty block between them)
            → merge the second into the first, separated by the style separator.

        The corresponding document paragraphs are merged in the XML tree and the
        block list is updated to match, so downstream steps see one block and one
        paragraph for each joined heading.
        """
        blocks = doc_structure.get("blocks", [])
        if not blocks:
            return

        separator = self._build_chapter_separator()

        # Build a fast lookup: stripped text → list of paragraph objects.
        # A list is used because duplicate text is possible in real documents.
        text_to_paras: dict[str, list] = {}
        for p in doc.paragraphs:
            key = p.text.strip()
            if key:
                text_to_paras.setdefault(key, []).append(p)

        i = 0
        while i < len(blocks):
            block = blocks[i]
            if block.get("type") != "heading_1":
                i += 1
                continue

            # Look ahead for the immediately following non-empty block
            next_content_idx = -1
            for j in range(i + 1, len(blocks)):
                if blocks[j].get("content", "").strip():
                    next_content_idx = j
                    break

            # Only merge when the very next non-empty block is also heading_1
            if next_content_idx == -1 or blocks[next_content_idx].get("type") != "heading_1":
                i += 1
                continue

            label_text = block.get("content", "").strip()
            title_text = blocks[next_content_idx].get("content", "").strip()

            label_paras = text_to_paras.get(label_text, [])
            title_paras = text_to_paras.get(title_text, [])

            if not label_paras or not title_paras:
                i += 1
                continue

            label_p = label_paras[0]
            title_p = title_paras[0]

            # Merge: append the title text to the label paragraph, then remove
            # the title paragraph from the XML tree entirely.
            merged_text = label_text + separator + title_text
            label_p.text = merged_text

            title_elem = title_p._element
            parent = title_elem.getparent()
            if parent is not None:
                parent.remove(title_elem)

            # Keep the text_to_paras index consistent for subsequent iterations
            text_to_paras[label_text] = [label_p]
            if len(title_paras) == 1:
                del text_to_paras[title_text]
            else:
                text_to_paras[title_text] = title_paras[1:]

            # Collapse the two blocks into one in the structure
            block["content"] = merged_text
            blocks.pop(next_content_idx)

            print(f"Merged headings: '{merged_text}'")
            # Do NOT increment i — re-check this position for back-to-back merges

    # -------------------------------------------------------------------------
    # Content formatting
    # -------------------------------------------------------------------------

    def _format_content_in_place(self, doc, doc_structure, has_title_page, title_page_boundary):
        """
        Format document content based on AI-detected structure.

        Performance notes
        -----------------
        * `doc.paragraphs` is a live XML walk — O(n) every call.  We materialise
          it once into `para_list` and never call `doc.paragraphs` again inside
          the loop.
        * Splits insert the new paragraph object directly into `para_list` at the
          correct position so the loop can visit it on the next iteration without
          missing or double-visiting anything.
        """
        # Build style map: stripped_text → (style_name, list_id)
        style_map: dict[str, tuple[str, any]] = {}
        for block in doc_structure.get("blocks", []):
            content = block.get("content", "").strip()
            block_type = block.get("type", "Normal")
            list_id = block.get("list_id") or block.get("attributes", {}).get("list_id")
            if content:
                style_map[content] = (STYLE_MAPPINGS.get(block_type, "Normal"), list_id)

        # Materialise paragraph list once — never re-query doc.paragraphs in this method
        para_list = list(doc.paragraphs)

        i = 0
        while i < len(para_list):
            p = para_list[i]
            current_text = p.text.strip()

            if not current_text:
                i += 1
                continue

            applied_style = False

            # --- Exact match ---
            if current_text in style_map:
                target_style_name, target_list_id = style_map[current_text]

                if target_style_name in doc.styles:
                    p.style = doc.styles[target_style_name]
                    applied_style = True
                else:
                    fallback_found = False
                    for suffix in [" 1", "1", " 2", "2"]:
                        variant_name = f"{target_style_name}{suffix}"
                        if variant_name in doc.styles:
                            p.style = doc.styles[variant_name]
                            target_style_name = variant_name
                            applied_style = True
                            fallback_found = True
                            print(f"Style '{target_style_name}' not found. Using variant '{variant_name}'.")
                            break

                    if not fallback_found:
                        print(f"Warning: Required style '{target_style_name}' not found. Falling back to Normal.")
                        p.style = doc.styles["Normal"]
                        applied_style = True

                if applied_style and target_style_name != "Normal":
                    if target_list_id is not None:
                        try:
                            p.list_id = target_list_id
                        except Exception:
                            pass
                    else:
                        # Remove residual list formatting when switching to a non-list style
                        if p._element.pPr is not None and not (
                            target_style_name.startswith("Heading") or "Appendix" in target_style_name
                        ):
                            numPr = p._element.pPr.find(qn('w:numPr'))
                            if numPr is not None:
                                p._element.pPr.remove(numPr)

            # --- Split detection (paragraph contains two logical blocks) ---
            if not applied_style:
                potential_split_info = None
                for mapped_text, (mapped_style, mapped_list_id) in style_map.items():
                    if (
                        mapped_text
                        and current_text.startswith(mapped_text)
                        and len(current_text) > len(mapped_text)
                    ):
                        remaining_part_raw = current_text[len(mapped_text):]
                        if remaining_part_raw.strip():
                            potential_split_info = (mapped_text, mapped_style, mapped_list_id)
                            break

                if potential_split_info:
                    mapped_text, mapped_style, mapped_list_id = potential_split_info
                    split_offset = len(mapped_text)

                    new_para = self._split_paragraph_at_offset(p, split_offset)

                    if new_para is not None:
                        if mapped_style in doc.styles:
                            new_para.style = doc.styles[mapped_style]
                        else:
                            new_para.style = doc.styles["Normal"]

                        if mapped_list_id is not None:
                            try:
                                new_para.list_id = mapped_list_id
                            except Exception:
                                pass

                        p.style = doc.styles["Normal"]
                        if p._element.pPr is not None and not (
                            mapped_style.startswith("Heading") or "Appendix" in mapped_style
                        ):
                            numPr = p._element.pPr.find(qn('w:numPr'))
                            if numPr is not None:
                                p._element.pPr.remove(numPr)

                        # Strip leading whitespace from the tail paragraph
                        found_content = False
                        for run in p.runs:
                            if found_content:
                                break
                            text = run.text
                            if not text:
                                continue
                            if text.strip() == "":
                                run.text = ""
                            else:
                                run.text = text.lstrip()
                                found_content = True

                        remaining_text = p.text.strip()

                        # Update para_list in-place: replace current slot with the
                        # new (prefix) paragraph, insert the tail right after.
                        # This ensures the tail is visited on the next iteration.
                        para_list[i] = new_para
                        para_list.insert(i + 1, p)

                        print(f"Split paragraph: '{mapped_text}' | '{remaining_text}'")
                        applied_style = True

            i += 1

    def _split_paragraph_at_offset(self, paragraph, offset: int):
        """
        Split a paragraph at the specified character offset, preserving formatting.
        Returns the NEW paragraph (which contains content BEFORE the offset).
        The original paragraph retains content AFTER the offset.
        """
        if offset < 0:
            raise ValueError("Offset cannot be negative")

        full_text = paragraph.text
        if offset == 0:
            return None

        current_pos = 0
        split_run_index = -1
        split_in_run_offset = 0

        for idx, run in enumerate(paragraph.runs):
            run_len = len(run.text)
            if current_pos <= offset < current_pos + run_len:
                split_run_index = idx
                split_in_run_offset = offset - current_pos
                break
            current_pos += run_len

        if split_run_index == -1 and offset == len(full_text):
            split_run_index = len(paragraph.runs)
            split_in_run_offset = 0
        elif split_run_index == -1:
            return None

        new_paragraph = paragraph.insert_paragraph_before('')

        all_runs = list(paragraph.runs)

        for idx, run in enumerate(all_runs):
            if idx < split_run_index:
                new_run = new_paragraph.add_run(run.text)
                self._copy_run_formatting(run, new_run)
                run.text = ""
            elif idx == split_run_index:
                original_text = run.text
                text_before = original_text[:split_in_run_offset]
                text_after = original_text[split_in_run_offset:]

                if text_before:
                    new_run = new_paragraph.add_run(text_before)
                    self._copy_run_formatting(run, new_run)

                run.text = text_after
            # Runs after the split point stay in the original paragraph unchanged

        return new_paragraph
        
    def _copy_run_formatting(self, source_run, target_run):
        """Copy formatting from source_run to target_run."""
        if not source_run or not target_run:
            return
            
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        
        if not hasattr(source_run, 'font') or not hasattr(target_run, 'font'):
            return
            
        source_font = source_run.font
        target_font = target_run.font
        
        if hasattr(source_font, 'name'):
            target_font.name = source_font.name
        if hasattr(source_font, 'size'):
            target_font.size = source_font.size

        if (
            hasattr(source_font, 'color')
            and source_font.color
            and source_font.color.type is not None  # has an explicit value
            and source_font.color.rgb
        ):
            target_font.color.rgb = source_font.color.rgb
                
        font_properties = [
            'highlight_color', 'subscript', 'superscript', 'strike',
            'double_strike', 'shadow', 'outline', 'rtl', 'imprint',
            'cs_bold', 'complex_script', 'hidden'
        ]
        
        for prop in font_properties:
            if hasattr(source_font, prop):
                try:
                    setattr(target_font, prop, getattr(source_font, prop))
                except (AttributeError, KeyError):
                    continue

    # -------------------------------------------------------------------------
    # Reference formatting
    # -------------------------------------------------------------------------

    def _format_references(self, doc, doc_structure):
        """
        Apply paragraph-level reference styles and sort references alphabetically.

        Run-level formatting (bold, italic, underline) is fully preserved because
        we only reassign the paragraph *style* and sort by swapping run XML subtrees
        between paragraphs rather than overwriting `p.text`.
        """
        ai_references = [
            block for block in doc_structure.get("blocks", [])
            if block.get("type") == "reference_list_item"
        ]

        if not ai_references:
            return

        # Collect expected reference texts (preserving order)
        reference_texts = [
            block.get("content", "").strip()
            for block in ai_references
            if block.get("content", "").strip()
        ]

        if not reference_texts:
            return

        # Map stripped text → paragraph object
        text_to_paragraph: dict[str, any] = {}
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                text_to_paragraph[text] = p

        reference_paragraphs = [
            text_to_paragraph[t] for t in reference_texts if t in text_to_paragraph
        ]

        if not reference_paragraphs:
            return

        # --- Apply style and paragraph-level formatting first ---
        style_config = self.selected_style_guide["styles"].get("References", {})

        for paragraph in reference_paragraphs:
            paragraph.style = doc.styles["References"]

            if "paragraph" in style_config:
                para_config = style_config["paragraph"]
                pf = paragraph.paragraph_format
                if "left_indent" in para_config:
                    pf.left_indent = para_config["left_indent"]
                if "first_line_indent" in para_config:
                    pf.first_line_indent = para_config["first_line_indent"]
                if "space_before" in para_config:
                    pf.space_before = para_config["space_before"]
                if "space_after" in para_config:
                    pf.space_after = para_config["space_after"]
                if "line_spacing" in para_config:
                    pf.line_spacing_rule = para_config["line_spacing"]

        # --- Sort by exchanging run content between paragraphs ---
        # We sort by a normalised sort key derived from the plain text, but transfer
        # the full XML run children so that all inline formatting survives intact.
        if len(reference_paragraphs) < 2:
            return

        def _sort_key(p):
            """Alphabetic key, ignoring leading numbering/brackets."""
            text = p.text.strip()
            cleaned = re.sub(r'^\s*[\(\[]?\d+[\)\]]?\.\?\s*', '', text)
            return cleaned.lower()

        # Collect run XML snapshots in current document order
        def _extract_runs_xml(paragraph) -> list:
            """Return a list of cloned run XML elements from a paragraph."""
            return [
                run._element.__class__(run._element.xml)
                for run in paragraph.runs
            ]

        # Build a parallel list of (sort_key, run_xml_snapshot) tuples
        snapshots = [(_sort_key(p), _extract_runs_xml(p)) for p in reference_paragraphs]
        snapshots.sort(key=lambda x: x[0])

        # Check whether sorting actually changes the order before touching the DOM
        original_keys = [s[0] for s in [(_sort_key(p), None) for p in reference_paragraphs]]
        sorted_keys   = [s[0] for s in snapshots]

        if original_keys == sorted_keys:
            return  # Already sorted — nothing to do

        # Replace run children in each paragraph with the sorted snapshot.
        # We do NOT touch paragraph properties (pPr) — only the runs (w:r elements).
        for paragraph, (_, run_xmls) in zip(reference_paragraphs, snapshots):
            p_elem = paragraph._p

            # Remove existing run elements
            for r_elem in p_elem.findall(qn('w:r')):
                p_elem.remove(r_elem)

            # Re-insert sorted run elements
            # Runs must be appended after the pPr element (paragraph properties), if present
            pPr_elem = p_elem.find(qn('w:pPr'))
            insert_after = pPr_elem if pPr_elem is not None else None

            for run_xml in run_xmls:
                if insert_after is not None:
                    insert_after.addnext(run_xml)
                    insert_after = run_xml
                else:
                    p_elem.append(run_xml)

    def _apply_font_properties(self, doc):
        """
        Explicitly applies required font properties (name, color, bold, size) to runs
        within paragraphs belonging to a specific set of target styles.
        """
        styles_config = self.selected_style_guide["styles"]
        
        target_style_names = {
            'Normal', 'Abstract', 'Title', 'Appendix Title',
            'Appendices Title', 'List Bullet', 'List Number', 'List Item',
            'Abstract Heading', 'Acknowledgement Heading', 'Acknowledgement Body',
            'Dedication Heading', 'Dedication Body', 'Preface Heading',
            'Preface Body', 'References Heading'
        }

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            
            if style_name in target_style_names:
                style_config = styles_config.get(style_name)
                
                if style_config and "font" in style_config:
                    font_config = style_config["font"]
                    
                    for run in paragraph.runs:
                        if "name" in font_config:
                            run.font.name = font_config["name"]
                        if "color" in font_config:
                            run.font.color.rgb = font_config["color"]
                        if "bold" in font_config:
                            run.font.bold = font_config["bold"]
                        
                        if style_name in {
                            'Normal', 'Abstract', 'Abstract Heading', 
                            'Acknowledgement Heading', 'Acknowledgement Body',
                            'Dedication Heading', 'Dedication Body', 'Preface Heading',
                            'Preface Body', 'References Heading'
                        } and "size" in font_config:
                            run.font.size = font_config["size"]

    def _apply_explicit_paragraph_properties(self, doc):
        """
        Explicitly applies all essential paragraph formatting properties (spacing,
        indentation, alignment, and keep options) from the style guide to each
        paragraph's format object.
        """
        styles_config = self.selected_style_guide["styles"]
        
        has_master_appendices = any(p.style.name == "Appendices Title" for p in doc.paragraphs)

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            
            if isinstance(style_name, str) and style_name.lower() == 'ds-markdown-paragraph':
                if 'Normal' in doc.styles:
                    paragraph.style = doc.styles['Normal']
                    style_name = 'Normal'
                    print(f"Converted ds-markdown-paragraph to 'Normal' for paragraph: '{paragraph.text[:60]}...'")
            
            if style_name not in styles_config:
                continue
                
            style_config = styles_config[style_name]
            if not style_config or "paragraph" not in style_config:
                continue
                
            para_config = style_config["paragraph"]
            para_format = paragraph.paragraph_format
            
            if "alignment" in para_config:
                alignment = para_config["alignment"]
                if style_name == "Appendix Title" and has_master_appendices:
                    alignment = WD_ALIGN_PARAGRAPH.LEFT
                para_format.alignment = alignment
            if "left_indent" in para_config:
                para_format.left_indent = para_config["left_indent"]
            if "right_indent" in para_config:
                para_format.right_indent = para_config["right_indent"]
            if "first_line_indent" in para_config:
                para_format.first_line_indent = para_config["first_line_indent"]

            if "space_before" in para_config:
                para_format.space_before = para_config["space_before"]
            if "space_after" in para_config:
                para_format.space_after = para_config["space_after"]
            if "line_spacing" in para_config:
                para_format.line_spacing = para_config["line_spacing"]

            if "keep_together" in para_config:
                para_format.keep_together = para_config["keep_together"]
            if "keep_with_next" in para_config:
                para_format.keep_with_next = para_config["keep_with_next"]
            if "page_break_before" in para_config:
                para_format.page_break_before = para_config["page_break_before"]
            if "widow_control" in para_config:
                para_format.widow_control = para_config["widow_control"]
            if "orphan_control" in para_config:
                try:
                    para_format.orphan_control = para_config["orphan_control"]
                except AttributeError:
                    pass
            
            if "outline_level" in para_config:
                try:
                    outline_level = para_config["outline_level"]
                    if style_name == "Appendix Title" and has_master_appendices:
                        outline_level = 1
                    para_format.outline_level = outline_level
                except AttributeError:
                    pass


    def _create_abstract_num(self, doc, format_type, lvl_text):
        """Create abstractNum element for list format."""
        try:
            numbering_part = doc.part.numbering_part
        except (NotImplementedError, KeyError, AttributeError):
            numbering_part = None
            
        if numbering_part is None:
            return None
        
        next_id = max([int(a.get(qn('w:abstractNumId'))) 
                       for a in numbering_part.element.findall(qn('w:abstractNum'))], default=0) + 1
        
        abstract = OxmlElement('w:abstractNum')
        abstract.set(qn('w:abstractNumId'), str(next_id))
        
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        
        fmt = OxmlElement('w:numFmt')
        fmt.set(qn('w:val'), format_type)
        lvl.append(fmt)
        
        text = OxmlElement('w:lvlText')
        text.set(qn('w:val'), lvl_text)
        lvl.append(text)
        
        if format_type == 'bullet':
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Symbol')
            rFonts.set(qn('w:hAnsi'), 'Symbol')
            rFonts.set(qn('w:hint'), 'default')
            rPr.append(rFonts)
            lvl.append(rPr)
            
            lvl_pPr = OxmlElement('w:pPr')
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '360')
            ind.set(qn('w:hanging'), '360')
            lvl_pPr.append(ind)
            lvl.append(lvl_pPr)

        abstract.append(lvl)
        numbering_part.element.insert(0, abstract)
        
        return next_id

    def _apply_list_properties(self, doc):
        """
        Apply list (bullet/numbered/alphabet) properties and handle list restart logic.
        """
        styles_config = self.selected_style_guide["styles"]
        try:
            numbering_part = doc.part.numbering_part
        except (NotImplementedError, KeyError, AttributeError):
            numbering_part = None
        if numbering_part is None:
            return
        
        def create_num(abstract_id):
            next_num_id = max([int(n.get(qn('w:numId'))) 
                              for n in numbering_part.element.findall(qn('w:num'))], default=5000) + 1
            num = OxmlElement('w:num')
            num.set(qn('w:numId'), str(next_num_id))
            abstract_ref = OxmlElement('w:abstractNumId')
            abstract_ref.set(qn('w:val'), str(abstract_id))
            num.append(abstract_ref)
            numbering_part.element.append(num)
            return next_num_id

        list_id_map = {} 
        
        next_decimal_id = 5000 
        next_alphabet_id = 7000

        previous_style = None
        
        last_used_decimal_id = next_decimal_id
        last_used_alphabet_id = next_alphabet_id
        last_used_bullet_id = None

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            
            is_list_item = style_name in {"List Bullet", "List Number", "List Item", "List Alphabet"}
            ai_list_id = paragraph.list_id if hasattr(paragraph, 'list_id') else None

            if not is_list_item:
                previous_style = style_name
                continue

            numbering_type = None
            target_num_id = None

            if style_name == "List Number":
                numbering_type = 'decimal'
                
                if ai_list_id is not None:
                    if ai_list_id in list_id_map:
                        target_num_id = list_id_map[ai_list_id]
                    else:
                        next_decimal_id += 1
                        target_num_id = next_decimal_id
                        list_id_map[ai_list_id] = target_num_id
                else:
                    should_restart = (previous_style != style_name)
                    if should_restart:
                        next_decimal_id += 1
                        target_num_id = next_decimal_id
                        last_used_decimal_id = target_num_id
                    else:
                        target_num_id = last_used_decimal_id

            elif style_name == "List Alphabet":
                numbering_type = 'lowerLetter'
                
                if ai_list_id is not None:
                    if ai_list_id in list_id_map:
                        target_num_id = list_id_map[ai_list_id]
                    else:
                        new_abstract_id = self._create_abstract_num(doc, 'lowerLetter', '%1)')
                        target_num_id = create_num(new_abstract_id)
                        list_id_map[ai_list_id] = target_num_id
                else:
                    should_restart = (previous_style != style_name)
                    if should_restart:
                        new_abstract_id = self._create_abstract_num(doc, 'lowerLetter', '%1)')
                        target_num_id = create_num(new_abstract_id)
                        last_used_alphabet_id = target_num_id
                    else:
                        target_num_id = last_used_alphabet_id

            elif style_name == "List Bullet":
                numbering_type = 'bullet'
                if ai_list_id is not None:
                    if ai_list_id in list_id_map:
                        target_num_id = list_id_map[ai_list_id]
                    else:
                        new_abstract_id = self._create_abstract_num(doc, 'bullet', '\uf0b7')
                        target_num_id = create_num(new_abstract_id)
                        list_id_map[ai_list_id] = target_num_id
                else:
                    if last_used_bullet_id is None or (previous_style != style_name):
                        new_abstract_id = self._create_abstract_num(doc, 'bullet', '\uf0b7')
                        target_num_id = create_num(new_abstract_id)
                        last_used_bullet_id = target_num_id
                    else:
                        target_num_id = last_used_bullet_id

            previous_style = style_name

            if style_name not in styles_config:
                continue

            p = paragraph._p
            
            if (numbering_type in {'decimal', 'lowerLetter', 'bullet'}) and target_num_id is not None:
                pPr = p.get_or_add_pPr()
                numPr = pPr.find(qn('w:numPr'))
                if numPr is None:
                    numPr = OxmlElement('w:numPr')
                    pPr.append(numPr)

                ilvl = numPr.find(qn('w:ilvl'))
                if ilvl is None:
                    ilvl = OxmlElement('w:ilvl')
                    numPr.append(ilvl)
                ilvl.set(qn('w:val'), '0')

                numId = numPr.find(qn('w:numId'))
                if numId is None:
                    numId = OxmlElement('w:numId')
                    numPr.append(numId)
                numId.set(qn('w:val'), str(target_num_id))

            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

    def _ensure_numbering_part(self, doc):
        try:
            return doc.part.numbering_part
        except (NotImplementedError, KeyError, AttributeError):
            return None
    
    def _apply_heading_styles(self, doc):
        """Apply heading font styles from the active style guide."""
        styles_config = self.selected_style_guide["styles"]

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name

            if not (style_name.startswith("Heading") or style_name in {"Appendix Title", "Appendices Title"}):
                continue

            if style_name.startswith("Heading"):
                try:
                    heading_level = int(style_name.split(' ')[1])
                except (IndexError, ValueError):
                    continue
                target_style_key = f"Heading {heading_level}"
            else:
                target_style_key = style_name

            style_entry = styles_config.get(target_style_key)
            if not style_entry:
                print(f"Warning: No style config found for '{target_style_key}'. Skipping heading font application.")
                continue

            font_config = style_entry.get("font", {})
            if not font_config:
                continue

            for run in paragraph.runs:
                if "name" in font_config:
                    run.font.name = font_config["name"]
                if "size" in font_config:
                    run.font.size = font_config["size"]
                if "bold" in font_config:
                    run.font.bold = font_config["bold"]
                if "italic" in font_config:
                    run.font.italic = font_config["italic"]
                if "color" in font_config:
                    run.font.color.rgb = font_config["color"]
                if "underline" in font_config:
                    run.font.underline = font_config["underline"]

    def _remove_blank_lines(self, doc, boundary_idx=0):
        """Remove empty paragraphs after the boundary index."""
        to_remove = [p for p in doc.paragraphs[boundary_idx:] 
                     if not p.text.strip() and not p.runs]
        
        for p in reversed(to_remove):
            p._element.getparent().remove(p._element)